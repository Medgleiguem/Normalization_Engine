"""
Report generator - creates academic-standard normalization reports
"""
from datetime import datetime
from typing import List
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.models.analysis_result import AnalysisResult, NormalizationStep
from app.models.table_model import Table

class ReportGenerator:
    """Generate comprehensive normalization analysis reports"""
    
    def __init__(self, analysis_result: AnalysisResult):
        self.analysis_result = analysis_result
        self.doc = Document()
        self._setup_styles()
    
    def generate_report(self, output_path: str) -> str:
        """Generate complete report and save to file"""
        # Title page
        self._add_title_page()
        
        # Table of contents (placeholder)
        self._add_section_break()
        
        # Executive summary
        self._add_executive_summary()
        
        # Original table analysis
        self._add_original_table_analysis()
        
        # Normalization process (step by step)
        self._add_normalization_process()
        
        # Final schema
        self._add_final_schema()
        
        # Theoretical background
        self._add_theoretical_background()
        
        # Conclusion
        self._add_conclusion()
        
        # References
        self._add_references()
        
        # Save document
        self.doc.save(output_path)
        return output_path
    
    def _setup_styles(self):
        """Configure document styles"""
        styles = self.doc.styles
        
        # Configure heading styles
        for i in range(1, 4):
            style = styles[f'Heading {i}']
            style.font.color.rgb = RGBColor(0, 51, 102)
    
    def _add_title_page(self):
        """Add title page"""
        # Title
        title = self.doc.add_heading('Database Normalization Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = self.doc.add_paragraph()
        subtitle.add_run(f'Analysis of: {self.analysis_result.original_table.name}').bold = True
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        
        # Metadata
        meta = self.doc.add_paragraph()
        meta.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y")}\n')
        meta.add_run(f'Analysis ID: {self.analysis_result.analysis_id or "N/A"}\n')
        meta.add_run(f'Final Normal Form: {self.analysis_result.normalization_steps[-1].to_nf.value if self.analysis_result.normalization_steps else "N/A"}')
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self._add_section_break()
    
    def _add_executive_summary(self):
        """Add executive summary section"""
        self.doc.add_heading('Executive Summary', 1)
        
        summary_text = f"""
This report presents a comprehensive analysis of the database normalization process applied to the 
'{self.analysis_result.original_table.name}' dataset. The analysis identified the original table's 
compliance with normal forms and systematically applied normalization techniques to achieve 
{self.analysis_result.normalization_steps[-1].to_nf.value if self.analysis_result.normalization_steps else "optimal normalization"}.

Key Findings:
• Original Normal Form: {self.analysis_result.current_normal_form.value}
• Target Normal Form: {self.analysis_result.target_normal_form.value}
• Normalization Steps Applied: {len(self.analysis_result.normalization_steps)}
• Final Number of Tables: {len(self.analysis_result.final_tables)}
• Total Violations Resolved: {len(self.analysis_result.get_all_violations())}

The normalization process successfully eliminated data redundancy, improved data integrity, 
and optimized the database structure for efficient querying and maintenance.
        """
        
        self.doc.add_paragraph(summary_text.strip())
        self._add_section_break()
    
    def _add_original_table_analysis(self):
        """Add original table analysis section"""
        self.doc.add_heading('Original Table Analysis', 1)
        
        table = self.analysis_result.original_table
        
        self.doc.add_heading('Table Structure', 2)
        self.doc.add_paragraph(f'Table Name: {table.name}')
        self.doc.add_paragraph(f'Number of Columns: {len(table.columns)}')
        self.doc.add_paragraph(f'Number of Rows (sample): {len(table.data)}')
        
        # Column details
        self.doc.add_heading('Column Specifications', 2)
        
        # Create table for columns
        col_table = self.doc.add_table(rows=1, cols=4)
        col_table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = col_table.rows[0].cells
        header_cells[0].text = 'Column Name'
        header_cells[1].text = 'Data Type'
        header_cells[2].text = 'Nullable'
        header_cells[3].text = 'Unique'
        
        # Data rows
        for col in table.columns:
            row_cells = col_table.add_row().cells
            row_cells[0].text = col.name
            row_cells[1].text = col.data_type.value
            row_cells[2].text = 'Yes' if col.nullable else 'No'
            row_cells[3].text = 'Yes' if col.unique else 'No'
        
        self.doc.add_paragraph()
        
        # Dependencies detected
        self.doc.add_heading('Dependencies Detected', 2)
        
        if table.functional_dependencies:
            self.doc.add_paragraph('Functional Dependencies:')
            for fd in table.functional_dependencies:
                self.doc.add_paragraph(f'  • {fd}', style='List Bullet')
        else:
            self.doc.add_paragraph('No functional dependencies detected.')
        
        if table.multi_valued_dependencies:
            self.doc.add_paragraph('Multi-Valued Dependencies:')
            for mvd in table.multi_valued_dependencies:
                self.doc.add_paragraph(f'  • {mvd}', style='List Bullet')
        
        # Primary key
        if table.primary_key:
            self.doc.add_paragraph(f'\nPrimary Key: {{{", ".join(sorted(table.primary_key))}}}')
        
        self._add_section_break()
    
    def _add_normalization_process(self):
        """Add detailed normalization process section"""
        self.doc.add_heading('Normalization Process', 1)
        
        for i, step in enumerate(self.analysis_result.normalization_steps, 1):
            self._add_normalization_step(step, i)
    
    def _add_normalization_step(self, step: NormalizationStep, step_number: int):
        """Add a single normalization step"""
        self.doc.add_heading(f'Step {step_number}: {step.from_nf.value} → {step.to_nf.value}', 2)
        
        # Explanation
        self.doc.add_paragraph(step.explanation.strip())
        
        # Violations
        if step.violations_found:
            self.doc.add_heading('Violations Identified', 3)
            for violation in step.violations_found:
                self.doc.add_paragraph(f'• {violation.description}', style='List Bullet')
                p = self.doc.add_paragraph()
                p.add_run('  Explanation: ').bold = True
                p.add_run(violation.explanation)
                p = self.doc.add_paragraph()
                p.add_run('  Resolution: ').bold = True
                p.add_run(violation.resolution)
                self.doc.add_paragraph()
        else:
            self.doc.add_paragraph('✓ No violations found. Table already complies with this normal form.')
        
        # Tables created
        if len(step.tables_created) > 1:
            self.doc.add_heading('Tables Created', 3)
            for table in step.tables_created:
                self.doc.add_paragraph(f'• {table.name}: {len(table.columns)} columns', style='List Bullet')
        
        self.doc.add_paragraph()
    
    def _add_final_schema(self):
        """Add final normalized schema section"""
        self.doc.add_heading('Final Normalized Schema', 1)
        
        self.doc.add_paragraph(f'The normalization process resulted in {len(self.analysis_result.final_tables)} table(s):')
        self.doc.add_paragraph()
        
        for table in self.analysis_result.final_tables:
            self.doc.add_heading(f'Table: {table.name}', 2)
            
            # Columns
            self.doc.add_paragraph('Columns:')
            for col in table.columns:
                self.doc.add_paragraph(f'  • {col.name} ({col.data_type.value})', style='List Bullet')
            
            # Primary key
            if table.primary_key:
                p = self.doc.add_paragraph()
                p.add_run('Primary Key: ').bold = True
                p.add_run(f'{{{", ".join(sorted(table.primary_key))}}}')
            
            # Foreign keys
            if table.foreign_keys:
                p = self.doc.add_paragraph()
                p.add_run('Foreign Keys:').bold = True
                for fk_col, (ref_table, ref_col) in table.foreign_keys.items():
                    self.doc.add_paragraph(f'  • {fk_col} → {ref_table}({ref_col})', style='List Bullet')
            
            self.doc.add_paragraph()
        
        self._add_section_break()
    
    def _add_theoretical_background(self):
        """Add theoretical background section"""
        self.doc.add_heading('Theoretical Background', 1)
        
        theory_text = """
Database normalization is a systematic approach to organizing data in a relational database to 
reduce redundancy and improve data integrity. The process involves decomposing tables into smaller, 
well-structured tables without losing information.

Normal Forms:

1NF (First Normal Form):
• All column values must be atomic (indivisible)
• No repeating groups or arrays
• Each column contains values of a single type

2NF (Second Normal Form):
• Must satisfy 1NF
• No partial dependencies (all non-key attributes must depend on the entire primary key)
• Relevant only for tables with composite primary keys

3NF (Third Normal Form):
• Must satisfy 2NF
• No transitive dependencies (non-key attributes must not depend on other non-key attributes)
• All attributes must depend directly on the primary key

BCNF (Boyce-Codd Normal Form):
• Must satisfy 3NF
• Every determinant must be a candidate key
• Stricter version of 3NF

4NF (Fourth Normal Form):
• Must satisfy BCNF
• No multi-valued dependencies
• Independent facts stored in separate tables

5NF (Fifth Normal Form / Project-Join Normal Form):
• Must satisfy 4NF
• No join dependencies
• Table cannot be decomposed further without loss of information

Benefits of Normalization:
• Eliminates data redundancy
• Improves data integrity and consistency
• Reduces update, insertion, and deletion anomalies
• Optimizes storage efficiency
• Enhances query performance
• Simplifies database maintenance
        """
        
        self.doc.add_paragraph(theory_text.strip())
        self._add_section_break()
    
    def _add_conclusion(self):
        """Add conclusion section"""
        self.doc.add_heading('Conclusion', 1)
        
        conclusion_text = f"""
This analysis successfully normalized the '{self.analysis_result.original_table.name}' dataset from 
{self.analysis_result.current_normal_form.value} to {self.analysis_result.normalization_steps[-1].to_nf.value if self.analysis_result.normalization_steps else "optimal form"}. 
The process identified and resolved {len(self.analysis_result.get_all_violations())} normalization violations, 
resulting in a well-structured database schema comprising {len(self.analysis_result.final_tables)} table(s).

The normalized schema provides:
• Enhanced data integrity through elimination of redundancy
• Improved query performance via optimized table structures
• Reduced storage requirements
• Simplified maintenance and updates
• Compliance with database design best practices

The accompanying MySQL script implements this normalized schema with proper constraints, indexes, 
and documentation, ready for deployment in a production environment.
        """
        
        self.doc.add_paragraph(conclusion_text.strip())
        self._add_section_break()
    
    def _add_references(self):
        """Add references section"""
        self.doc.add_heading('References', 1)
        
        references = [
            'Codd, E. F. (1970). "A Relational Model of Data for Large Shared Data Banks". Communications of the ACM.',
            'Date, C. J. (2003). "An Introduction to Database Systems" (8th ed.). Addison-Wesley.',
            'Elmasri, R., & Navathe, S. B. (2015). "Fundamentals of Database Systems" (7th ed.). Pearson.',
            'Kent, W. (1983). "A Simple Guide to Five Normal Forms in Relational Database Theory". Communications of the ACM.',
            'Silberschatz, A., Korth, H. F., & Sudarshan, S. (2019). "Database System Concepts" (7th ed.). McGraw-Hill.'
        ]
        
        for i, ref in enumerate(references, 1):
            self.doc.add_paragraph(f'[{i}] {ref}')
    
    def _add_section_break(self):
        """Add a section break"""
        self.doc.add_page_break()
